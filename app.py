import streamlit as st
from openai import OpenAI
import json
import os
import pandas as pd
from io import BytesIO
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. 页面配置：专业商务风格
st.set_page_config(page_title="律盾 AI - 智慧法务旗舰版", page_icon="⚖️", layout="wide")

# 自定义 CSS 美化界面
st.markdown("""
    <style>
    .main { background-color: #f8f9fa; }
    .stTabs [data-baseweb="tab-list"] { gap: 24px; }
    .stTabs [data-baseweb="tab"] { height: 50px; white-space: pre-wrap; background-color: #ffffff; border-radius: 5px; border: 1px solid #eee; }
    .stTabs [aria-selected="true"] { background-color: #1E3A8A !important; color: white !important; }
    </style>
    """, unsafe_allow_html=True)

# --- 数据库持久化 ---
USER_DB_FILE = "users_data.json"

def load_users():
    if os.path.exists(USER_DB_FILE):
        try:
            with open(USER_DB_FILE, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return {"admin": {"password": "888", "role": "admin", "count": 999}}

def save_users(users_data):
    with open(USER_DB_FILE, "w", encoding="utf-8") as f:
        json.dump(users_data, f, ensure_ascii=False, indent=4)

if 'users' not in st.session_state:
    st.session_state.users = load_users()

# 2. API 配置
try:
    client = OpenAI(api_key=st.secrets["api_key"], base_url="https://api.deepseek.com")
except Exception as e:
    st.error("❌ API 配置失效，请检查后台 Secrets。")
    st.stop()

# --- 高级工具函数：生成专业 Word 报告 ---
def generate_pro_docx(title, content):
    doc = docx.Document()
    # 标题
    header = doc.add_heading(title, 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 正文
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            run = p.add_run()
            font = run.font
            font.name = '微软雅黑'
            font.size = Pt(11)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 3. 侧边栏：登录与管理
with st.sidebar:
    st.title("🛡️ 会员管理中心")
    if "logged_in" not in st.session_state: st.session_state.logged_in = False

    if not st.session_state.logged_in:
        t1, t2 = st.tabs(["🔑 登录", "📝 手机注册"])
        with t1:
            u = st.text_input("手机号", key="l_u")
            p = st.text_input("密码", type="password", key="l_p")
            if st.button("立即进入系统", use_container_width=True):
                if u in st.session_state.users and st.session_state.users[u]["password"] == p:
                    st.session_state.logged_in, st.session_state.user = True, u
                    st.session_state.role = st.session_state.users[u].get("role", "user")
                    st.rerun()
                else: st.error("❌ 验证失败")
        with t2:
            st.caption("🎁 新用户注册立赠 5 次额度")
            ru = st.text_input("📱 手机号", key="r_u")
            rp = st.text_input("🔒 设置密码", type="password", key="r_p")
            if st.button("完成注册并领取额度", use_container_width=True):
                if len(ru) == 11 and rp:
                    if ru not in st.session_state.users:
                        st.session_state.users[ru] = {"password": rp, "role": "user", "count": 5}
                        save_users(st.session_state.users)
                        st.success("✅ 注册成功！请登录")
                    else: st.error("❌ 账号已存在")
        st.stop()
    else:
        st.success(f"👤 已认证：{st.session_state.user}")
        curr_u = st.session_state.users[st.session_state.user]
        st.metric("剩余可用点数", f"{curr_u['count']} 次")
        
        if st.session_state.role == "admin":
            with st.expander("👑 超级管理员后台"):
                target = st.selectbox("选择账号", list(st.session_state.users.keys()))
                num = st.number_input("充值点数", 1, 1000, 50)
                if st.button("立即充值"):
                    st.session_state.users[target]["count"] += num
                    save_users(st.session_state.users)
                    st.success("充值成功")
                    st.rerun()

        st.divider()
        st.info("🆘 客服微信：**linwlang12**")
        if st.button("🚪 安全退出", use_container_width=True):
            st.session_state.logged_in = False
            st.rerun()

# 4. 主界面核心功能
st.title("⚖️ 律盾 AI - 专业私人律师旗舰版")

# 拦截与扣费逻辑
def check_and_use():
    if st.session_state.users[st.session_state.user]["count"] <= 0:
        st.error("⚠️ 余额不足！请添加客服微信 **linwlang12** 充值获取更多额度。")
        return False
    return True

tabs = st.tabs(["🔍 合同风险深度扫描", "✍️ 专业法律文书起草", "🌐 国际法律文本翻译"])

# --- 模块 1：高级风险分析 ---
with tabs[0]:
    st.subheader("🎯 风险评估与条款审查")
    col1, col2 = st.columns([3, 1])
    with col1:
        txt = st.text_area("请粘贴合同或条款内容：", height=300, placeholder="支持最高 5000 字扫描...", key="audit_in")
        char_count = len(txt)
        st.caption(f"当前字数：{char_count} / 5000")
    with col2:
        st.write("🔧 扫描偏好")
        opt = st.radio("深度等级", ["标准扫描", "深度挖掘（含法律依据）", "霸王条款专项"])
        is_en = st.checkbox("生成双语对照报告")

    if st.button("🚀 开始 AI 专家级审计", use_container_width=True):
        if char_count > 5000: st.error("❌ 文本过长，请分段处理")
        elif txt and check_and_use():
            with st.spinner("⚖️ 律师团队正在模拟审查，请稍候..."):
                prompt = f"你是一位资深中国执业律师。请针对以下内容执行【{opt}】分析。必须包含：1.风险评级 2.核心风险点 3.对应的法律法规依据（如民法典条文） 4.具体的修改建议。"
                if is_en: prompt += " 请先提供详细中文分析，再用 '---SPLIT---' 分隔，最后提供对应的英文摘要。"
                
                try:
                    res = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"system","content":prompt},{"role":"user","content":txt}])
                    ans = res.choices[0].message.content
                    
                    if "---SPLIT---" in ans:
                        cn, en = ans.split("---SPLIT---")
                        c_l, c_r = st.columns(2)
                        with c_l: st.error("🇨🇳 中文深度报告"); st.write(cn)
                        with c_r: st.warning("🇺🇸 English Report"); st.write(en)
                    else:
                        st.markdown("### 📜 法律审查报告")
                        st.info(ans)
                    
                    st.divider()
                    st.download_button("📥 下载专业版 Word 报告", data=generate_pro_docx("法律风险审计报告", ans.replace("---SPLIT---","\n")), file_name="律盾AI分析报告.docx")
                    
                    st.session_state.users[st.session_state.user]["count"] -= 1
                    save_users(st.session_state.users)
                except Exception as e: st.error(f"分析失败: {e}")

# --- 模块 2：专业起草 ---
with tabs[1]:
    st.subheader("✍️ 法律文书标准起草")
    req = st.text_area("起草需求描述：", placeholder="请详细描述：文书类型、涉及金额、双方责任、违约金要求等...", height=200)
    if st.button("✨ 生成法律文书草案", use_container_width=True):
        if req and check_and_use():
            with st.spinner("正在检索法律模板并起草..."):
                res = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"system","content":"你是一位严谨的中国律师。请起草一份格式规范、用词专业的法律文书。"},{"role":"user","content":req}])
                ans = res.choices[0].message.content
                st.markdown("### 📄 文书预览")
                st.code(ans, language="markdown")
                st.download_button("📥 下载 Word 文稿", data=generate_pro_docx("法律文书起草", ans), file_name="起草文书.docx")
                st.session_state.users[st.session_state.user]["count"] -= 1
                save_users(st.session_state.users)

# --- 模块 3：法律翻译 ---
with tabs[2]:
    st.subheader("🌐 国际法律术语级翻译")
    tin = st.text_area("待翻译文本：", height=250, placeholder="粘贴需要翻译的合同、判决书或法条...")
    if st.button("🛰️ 执行专业法律翻译", use_container_width=True):
        if tin and check_and_use():
            with st.spinner("正在进行术语核对与翻译..."):
                res = client.chat.completions.create(model="deepseek-chat", messages=[{"role":"system","content":"你是一位精通中英法律体系的翻译专家。请确保法律术语（如indemnify, Force Majeure等）使用准确且严谨。"},{"role":"user","content":tin}])
                ans = res.choices[0].message.content
                st.success("✅ 翻译完成")
                st.write(ans)
                st.download_button("📥 下载翻译件", data=generate_pro_docx("专业法律翻译报告", ans), file_name="翻译结果.docx")
                st.session_state.users[st.session_state.user]["count"] -= 1
                save_users(st.session_state.users)
